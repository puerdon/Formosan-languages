import re
import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from streamlit_gsheets import GSheetsConnection
import time

st.markdown("# 上傳語料")
# st.write(st.session_state['last_update_time'])
conn = st.connection("gsheets", type=GSheetsConnection)

# st.dataframe(conn.read(worksheet="last updated", ttl=0))
LANG_TABLE = {
"阿美語":["北部阿美語", "中部阿美語", "海岸阿美語", "馬蘭阿美語", "恆春阿美語"],
"泰雅語":["賽考利克泰雅語","澤敖利泰雅語","汶水泰雅語","萬大泰雅語","四季泰雅語","宜蘭澤敖利泰雅語"],
"排灣語":["東排灣語","北排灣語","中排灣語","南排灣語"],
"布農語":["卓群布農語","卡群布農語","丹群布農語","巒群布農語","郡群布農語"],
"卑南語":["南王卑南語","知本卑南語","初鹿卑南語","建和卑南語"],
"魯凱語":["東魯凱語","霧臺魯凱語","大武魯凱語","多納魯凱語","茂林魯凱語","萬山魯凱語"],
"鄒語":[],
"卡那卡那富語":[],
"拉阿魯哇語":[],
"賽夏語":[],
"雅美語":[],
"邵語":[],
"噶瑪蘭語":[],
"太魯閣語":[],
"撒奇萊雅語":[],
"賽德克語":["都達語","德固達雅語","德魯固語"]
}

LANG_ENG_TABLE = {
"阿美語":"Amis",
"泰雅語":"Atayal",
"排灣語":"Paiwan",
"布農語":"Bunun",
"卑南語":"Puyuma",
"魯凱語":"Rukai",
"鄒語":"Tsou",
"卡那卡那富語":"Kanakanavu",
"拉阿魯哇語":"Saaroa",
"賽夏語":"Saisiyat",
"雅美語":"Yami",
"邵語":"Thao",
"噶瑪蘭語":"Kavalan",
"太魯閣語":"Truku",
"撒奇萊雅語":"Sakizaya",
"賽德克語":"Seediq"
}

def parse_docx(file_path):
    # 读取 .docx 文件
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])+'\n'  # 合并所有段落为文本

    # 定义正则模式
    pattern = r"(A|RA|G|M|RM|E|RE):\s*(.*?)\n"
    matches = re.findall(pattern, text)
    
    # 按段落类型组织内容
    parsed_data = {}
    for key, value in matches:

        if key not in parsed_data:
            parsed_data[key] = []
        parsed_data[key].append(value.strip())
    
    return pd.DataFrame.from_dict(parsed_data)

# def 檢查各檔案的每個句子的tag(doc):

#     having_five_tags = False
#     after_title_part = False
#     after_new_line = False

#     all_sentences = []

#     for para in doc.paragraphs:
        
# #         print(f'{para.text}\n###')
        
#         ################################
#         # BEGIN: 處理語料前面的metadata #       
#         ################################
        
#         if not after_title_part:
#             if not para.text.startswith("A:") and\
#             not para.text.startswith("RA:") and\
#             not para.text.startswith("G:") and\
#             not para.text.startswith("M:") and\
#             not para.text.startswith("RM:"):
#                 continue

#         after_title_part = True
        
        

#         # 遇到換行
#         if len(para.text.strip()) == 0:
#             after_new_line = True
#             all_sentences.append(dict())

#         if not para.text.startswith("A:") and\
#         not para.text.startswith("RA:") and\
#         not para.text.startswith("G:") and\
#         not para.text.startswith("M:") and\
#         not para.text.startswith("RM:"):
#             continue

#         for p in para.text.split('\n'):
#             if p.startswith("A:") or p.startswith("RA:") or p.startswith("G:") or p.startswith("M:") or p.startswith("RM:"):
#                 tag = p.split(":")[0]
#                 content = p.lstrip(tag + ":").strip()
#                 if content == 'none':
#                     content = None

#                 if len(all_sentences) == 0:
#                     all_sentences.append(dict())

#                 all_sentences[-1][tag] = content

#     return all_sentences

uploaded_files = st.file_uploader("Choose a file", accept_multiple_files=True)

st.markdown("## 語料預覽")

# 初始化一个列表，用于存储所有 DataFrame
dataframes = []

if len(uploaded_files) > 0:
    for f in uploaded_files:
        # 将内容提取并显示
    
        # all_sentences = 檢查各檔案的每個句子的tag(doc)
        # all_sentences = pd.DataFrame(all_sentences)
        all_sentences = parse_docx(f)
        all_sentences = all_sentences.replace({np.nan: None})
        dataframes.append(all_sentences)
    
    combined_df = pd.concat(dataframes, ignore_index=True)
    st.dataframe(combined_df)

    ab_column_name = st.selectbox("族語欄位", options=combined_df.columns)
    trans_column_name = st.selectbox("翻譯欄位", options=combined_df.columns)
    

    ## TODO 要驗證是否每個欄位都有東西
    
    # for index, row in all_sentences.iterrows():
    #     ab = ""
    #     ch = ""
        
    #     try:
    #         ab = row['A'] if row['RA'] is None or len(row['RA']) == 0 else row['RA']
    #         ch = row['M'] if row['RM'] is None or len(row['RM']) == 0 else row['RM']
    #     except Exception as e:
    #         print(x)
    #         print(e)

    #     if ab is None or ch is None:
    #         continue
        
    #     df= pd.DataFrame({'Ab':[ab], 'Ch':[ch], 'Source': [f'{x.name[:4]}_{str(index)}']})
    
    #     results = pd.concat([results, df], ignore_index = True)
    
st.markdown("## 語料詳情")
lang_option = st.selectbox("語言別",list(LANG_TABLE))
dialect_option = st.selectbox("方言別", options=LANG_TABLE[lang_option])
dialect_option = dialect_option.rstrip(lang_option)
name = st.text_input("語料主來源標題")
sub_name = st.text_input("語料子標題")
password = st.text_input("上傳密碼")

# st.write("語言別:", lang_option)
# st.write("方言別:", dialect_option)
# st.write("來源:", name)
# st.write("子文件:", sub_name)

if st.button('update') and password == st.secrets["upload_pwd"]:
    conn = st.connection("gsheets", type=GSheetsConnection)
    original_df = conn.read(worksheet="user corpus", ttl=0)
    
    all_sentences = combined_df[[ab_column_name, trans_column_name]]
    all_sentences = all_sentences.rename({ab_column_name: "Ab", trans_column_name: "Ch"}, axis='columns')
    all_sentences['Lang_En'] = LANG_ENG_TABLE[lang_option]
    all_sentences['Lang_Ch'] = lang_option[:-1] + '_' + dialect_option if dialect_option != '' else lang_option[:-1]
    all_sentences['From'] = name
    all_sentences['Source'] = [sub_name + "_" + str(i) for i in range(1, len(all_sentences) + 1)]
    result_df = pd.concat([original_df, all_sentences], ignore_index=True)

    with st.spinner('上傳語料中...'):
        conn.update(worksheet="user corpus", data=result_df)
        last_update_time = conn.read(worksheet="last updated", ttl=0)
        current_time = {'time': [int(time.time())]}
        conn.update(worksheet="last updated", data=pd.DataFrame.from_dict(current_time))
    st.success('上傳成功!', icon="✅")

